from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Color palette ──────────────────────────────────────────────
NAVY   = RGBColor(0x14, 0x32, 0x78)   # 20,50,120
GRAY   = RGBColor(0x50, 0x50, 0x50)
LGRAY  = RGBColor(0x90, 0x90, 0x90)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GREEN  = RGBColor(0x14, 0x82, 0x3C)
ORANGE = RGBColor(0xC8, 0x8C, 0x00)
RED    = RGBColor(0xB4, 0x32, 0x32)
CALLOUT_BG = "E6F0FF"  # light blue hex for shading
WARN_BG    = "FFF8E6"
GOOD_BG    = "E6FFE6"

FONT_NAME = "맑은 고딕"


# ── Helpers ────────────────────────────────────────────────────

def set_font(run, size=10, bold=False, color=None, italic=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    # Force Korean font
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    rPr.insert(0, rFonts)


def set_para_font(para, size=10, bold=False, color=None):
    for run in para.runs:
        set_font(run, size, bold, color)


def add_heading(doc, text, level=1):
    """Custom heading with navy color and no extra space above."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run(text)
    size = {1: 15, 2: 12, 3: 11}.get(level, 10)
    set_font(run, size=size, bold=True, color=NAVY)
    if level == 1:
        # Bottom border
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "143278")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return para


def add_body(doc, text, bold_parts=None, indent=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(3)
    if indent:
        para.paragraph_format.left_indent = Cm(0.5)
    run = para.add_run(text)
    set_font(run, size=10, color=GRAY)
    return para


def add_bullet(doc, text, indent_level=1):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(2)
    para.paragraph_format.left_indent  = Cm(0.5 * indent_level)
    run = para.add_run(text)
    set_font(run, size=10, color=GRAY)
    return para


def add_numbered(doc, text, indent_level=1):
    para = doc.add_paragraph(style="List Number")
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(2)
    para.paragraph_format.left_indent  = Cm(0.5 * indent_level)
    run = para.add_run(text)
    set_font(run, size=10, color=GRAY)
    return para


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=9,
                  align=WD_ALIGN_PARAGRAPH.LEFT, bg=None):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    run = para.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    if bg:
        shade_cell(cell, bg)


def add_table(doc, headers, rows, col_widths=None, header_bg="143278",
              alt_bg="EEF2FF"):
    """Generic bordered table."""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Column widths
    total_cm = 16.0
    if col_widths:
        widths = col_widths
    else:
        widths = [total_cm / n_cols] * n_cols

    for i, w in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = Cm(w)

    # Header row
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h,
                      bold=True, color=WHITE, size=9,
                      align=WD_ALIGN_PARAGRAPH.CENTER, bg=header_bg)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        bg = alt_bg if r_idx % 2 == 1 else "FAFAFA"
        for c_idx, cell_text in enumerate(row_data):
            set_cell_text(table.rows[r_idx + 1].cells[c_idx],
                          cell_text, size=9, bg=bg)

    doc.add_paragraph()
    return table


def add_callout(doc, title, text, bg_hex=CALLOUT_BG):
    """Shaded callout box using a 1-column table."""
    table = doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    # Title row
    set_cell_text(table.rows[0].cells[0], f"  {title}",
                  bold=True, color=NAVY, size=10,
                  align=WD_ALIGN_PARAGRAPH.LEFT, bg=bg_hex)
    # Body row
    cell = table.rows[1].cells[0]
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    para.paragraph_format.left_indent  = Cm(0.3)
    run = para.add_run(text)
    set_font(run, size=10, color=GRAY)
    shade_cell(cell, bg_hex)
    doc.add_paragraph()
    return table


def add_icon_bullet(doc, icon, label, text, icon_color):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(3)
    para.paragraph_format.left_indent  = Cm(0.3)
    r_icon = para.add_run(icon + " ")
    set_font(r_icon, size=10, bold=True, color=icon_color)
    r_label = para.add_run(label + ": ")
    set_font(r_label, size=10, bold=True)
    r_text = para.add_run(text)
    set_font(r_text, size=10, color=GRAY)
    return para


# ══════════════════════════════════════════════════════════════
#  Build document
# ══════════════════════════════════════════════════════════════

def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Cover ─────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run("내부 검토 보고서")
    set_font(r, size=11, bold=True, color=NAVY)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run("AX Sprint Workshop")
    set_font(r2, size=26, bold=True, color=RGBColor(0x14, 0x1E, 0x3C))

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(20)
    r3 = p3.add_run("견적 종합 분석 보고서")
    set_font(r3, size=20, bold=True, color=NAVY)

    for line in [
        "제안사: 아리송컴퍼니 (대표 송락현)",
        "검토 대상: 올거나이즈(Allganize) Alli 플랫폼 교육 워크숍",
        "견적 총액: 14,000천원 (VAT 별도)",
        "",
        "작성일: 2026년 4월 12일  |  하나증권 인재개발실",
        "외부 배포 금지  |  내부 검토용",
    ]:
        pp = doc.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.paragraph_format.space_before = Pt(0)
        pp.paragraph_format.space_after  = Pt(3)
        rr = pp.add_run(line)
        set_font(rr, size=11, color=LGRAY)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 1. 견적 구조 요약
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "1. 견적 구조 요약", 1)
    add_body(doc, "아리송컴퍼니(대표: 송락현)가 제안한 AX Sprint Workshop 견적은 3개 항목으로 구성됩니다.")

    add_table(doc,
        headers=["항목", "금액(천원)", "포함 내용"],
        rows=[
            ["프로그램 기획 및 교안 설계", "4,000",
             "맞춤형 교안 · 그룹사 커리큘럼 자문 · 기술가이드 반영"],
            ["오프라인 워크숍 진행 (2회/6시간)", "6,000",
             "총 2회 (1회차 기획/설계, 2회차 구축/검증)"],
            ["기술 지원", "4,000",
             "올거나이즈 기반 케이스별 솔루션 테스트"],
            ["합계 (VAT 제외)", "14,000", ""],
            ["합계 (VAT 포함)", "15,400", ""],
        ],
        col_widths=[5.5, 3.0, 7.5],
    )

    add_heading(doc, "기타 조건", 3)
    add_bullet(doc, "SaaS 라이선스 비용: 하나증권 부담")
    add_bullet(doc, "산출 에이전트 소유권: 하나증권 귀속")
    add_bullet(doc, "AX Sprint 프레임워크 저작권: 아리송컴퍼니 귀속")

    # ══════════════════════════════════════════════════════════
    # 2. 견적 타당성 분석
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "2. 견적 타당성 분석 (시세 비교)", 1)

    add_heading(doc, "2-1. 공공기관 기준 대비 (2022 서울시 인재원 — 하한선 참고)", 2)
    add_body(doc, "공공기관 기준은 민간보다 낮아 직접 비교는 부적절하지만, 절대적 하한선 파악에 유용합니다.")
    add_table(doc,
        headers=["등급", "기본(1h)", "6시간 합계", "비고"],
        rows=[
            ["일반교육 특강급 (차관급 이상)", "400천원", "1,400천원", "2022년 기준"],
            ["일반교육 전문급 (5년 이상)", "240천원", "840천원", ""],
            ["퍼실리테이터 FT1급 (15년 이상)", "230천원", "980천원", ""],
            ["정보화교육 1급", "150천원", "650천원", ""],
            ["교육 자문", "100천원/h", "—", ""],
        ],
        col_widths=[5.5, 2.5, 3.0, 5.0],
    )
    add_callout(doc,
        "2026년 환산 시",
        "2022년 기준 물가 상승(연 3~4%) 반영 시 특강급 최고 등급도 6시간에 약 170만원 수준.\n"
        "이는 아리송컴퍼니 워크숍 진행비 600만원의 약 1/4 수준입니다.",
        bg_hex=WARN_BG)

    add_heading(doc, "2-2. 2026년 민간 AI 교육 시세", 2)
    add_bullet(doc, "IT/디지털 전환 분야 강사: 시간당 50~200만원")
    add_bullet(doc, "일반 민간 AI 실습 워크숍: 시간당 80~150만원 주류")
    add_bullet(doc, "실습형 워크숍은 강의형 대비 1.5~2배 높게 형성 (준비 공수 반영)")

    add_heading(doc, "2-3. 항목별 적정성 판단", 2)
    add_icon_bullet(doc, "▶", "워크숍 진행 600만원",
        "시간당 100만원 — 민간 시세 중간 수준. 강사 1인 투입 시 비쌈, 2인 투입 시 적정선",
        NAVY)
    add_icon_bullet(doc, "▶", "기획/교안 설계 400만원",
        "\"커리큘럼 자문\" 범위 불명확. 순수 교안만이면 과다 (통상 강의료의 50~100%)",
        NAVY)
    add_icon_bullet(doc, "▶", "기술 지원 400만원",
        "가장 근거가 약한 항목. 올거나이즈 공식 온보딩과 중복되며 인재개발실에서 대체 가능",
        NAVY)

    add_callout(doc,
        "종합 판단",
        "총 1,400만원 ÷ 6시간 = 시간당 약 233만원 (기획·지원 포함)\n"
        "민간 시세 상위권에 해당하며, 기획비와 기술지원비 각 400만원이 총액을 과도하게 끌어올린 구조.\n"
        "실제 시세 적정선: 450~800만원 (VAT 별도)",
        bg_hex="FFE0E0")

    # ══════════════════════════════════════════════════════════
    # 3. 불명확 항목
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "3. 견적 불명확 항목 — 반드시 확인해야 할 사항", 1)

    add_heading(doc, "3-1. 필수 확인 사항 (아리송컴퍼니에 서면 요청 권장)", 2)

    items = [
        ("투입 인력 계획 미제시",
         "강사/퍼실리테이터 몇 명이 투입되는지, 각 인력의 경력·자격 전혀 명시되지 않음. "
         "인일(Man-Day) 기반 산출 근거 없어 적정성 검증 불가."),
        ("\"기술 지원\" 400만원의 구체적 범위",
         "\"케이스별 솔루션 테스트\"가 몇 건인지, PoC급인지 프로토타입급인지, 결과 산출물이 무엇인지 불명확. "
         "실질적으로 강사의 올거나이즈 사전 학습 비용이 포함된 것으로 추정."),
        ("\"그룹사 AI 리터러시 커리큘럼 자문\"의 범위",
         "이 워크숍과 별개의 전사 컨설팅이라면 별도 항목으로 분리해야 함. "
         "자문 횟수·시간·산출물이 명시되어야 계약 가능."),
        ("교안/교재 소유권 모호",
         "\"AX Sprint 프레임워크 저작권은 아리송컴퍼니 귀속\"이라 명시되었으나, "
         "400만원 들여 만든 하나증권 맞춤 교안과 실습 자료의 재사용권 불명확. "
         "하나증권이 추후 자체 교육 시 이 교안 활용 가능 여부 확인 필요."),
        ("사후 지원 포함 여부",
         "워크숍 종료 후 Q&A, 추가 기술 지원, 결과물 고도화 지원 포함 여부 명시 없음."),
    ]
    for i, (label, text) in enumerate(items, 1):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after  = Pt(2)
        para.paragraph_format.left_indent  = Cm(0.3)
        r_num = para.add_run(f"{i}. ")
        set_font(r_num, size=10, bold=True, color=NAVY)
        r_label = para.add_run(label + " — ")
        set_font(r_label, size=10, bold=True)
        r_text = para.add_run(text)
        set_font(r_text, size=10, color=GRAY)

    add_heading(doc, "3-2. 추가 확인 권장 사항", 2)
    extra = [
        ("6. 아리송컴퍼니 레퍼런스/포트폴리오",
         "웹 검색에서 해당 업체 정보가 거의 확인되지 않음. 유사 워크숍 수행 실적과 올거나이즈 플랫폼 전문성 검증 필요."),
        ("7. 올거나이즈 공식 파트너 여부",
         "아리송컴퍼니가 올거나이즈의 공인 파트너/리셀러인지, 독립 교육업체인지 확인 필요. "
         "올거나이즈 자체 교육/온보딩 프로그램 제공 여부도 확인 (라이선스에 포함 가능성 있음)."),
        ("8. 20명 대상 실습 구성",
         "팀 구성(몇 팀?), 팀별 멘토 배치 여부, 올거나이즈 계정 구성(팀별 1개 vs 개인별) 등 실습 환경 세부 사항 미정."),
    ]
    for label, text in extra:
        add_icon_bullet(doc, "○", label, text, LGRAY)

    # ══════════════════════════════════════════════════════════
    # 4. 올거나이즈 온보딩
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "4. 올거나이즈(Allganize) 자체 온보딩 확인 결과", 1)

    add_callout(doc,
        "핵심 발견사항",
        "올거나이즈 공식 가이드 사이트(guide.allganize.ai)를 직접 확인한 결과:\n"
        "• 한국어 온보딩 가이드 전체가 무료 공개\n"
        "• 로그인 후 플랫폼 내 튜토리얼 자동 제공\n"
        "• 엔터프라이즈 계약 고객에게 담당 CS + 기술 온보딩 지원 기본 포함\n\n"
        "→ 아리송컴퍼니의 \"기술 지원\" 400만원 항목 대부분이 올거나이즈 자체 리소스와 중복됩니다.",
        bg_hex=GOOD_BG)

    add_heading(doc, "플랫폼 기술 난이도 등급", 2)
    add_table(doc,
        headers=["기능 영역", "난이도", "비고"],
        rows=[
            ["기본 사용 (문서 업로드, 간단 Q&A 에이전트)", "낮음", "비개발자도 2~3시간이면 기본 동작 확인"],
            ["LLM App Builder 워크플로우 구축", "낮음~중간", "드래그앤드롭; 논리적 사고력 있으면 가능"],
            ["RAG 최적화 (청킹 전략, 프롬프트 튜닝)", "중간", "문서 구조별 최적화 경험·반복 실험 필요"],
            ["API 연동 및 외부 시스템 통합", "중간~높음", "개발 지식 필요, 비개발자 독립 수행 어려움"],
            ["프로덕션 에이전트 배포·운영", "높음", "모니터링·오류 처리·성능 최적화 운영 노하우"],
        ],
        col_widths=[6.5, 2.5, 7.0],
    )

    add_heading(doc, "6시간 워크숍에서 현실적으로 달성 가능한 수준", 2)
    add_bullet(doc, "올거나이즈 기본 인터페이스 숙달")
    add_bullet(doc, "간단한 RAG 기반 문서 Q&A 에이전트 1개 구축")
    add_bullet(doc, "노드 3~5개 수준의 간단한 워크플로우 구현")
    add_callout(doc,
        "주의",
        "제안서의 \"작동하는 AI 에이전트(MVP) 직접 구현\"은 마케팅 표현입니다.\n"
        "6시간 내 가능한 것은 데모 수준의 프로토타입이며, 프로덕션 배포 수준과는 차이가 있습니다.",
        bg_hex=WARN_BG)

    # ══════════════════════════════════════════════════════════
    # 5. 인재개발실 커버 범위
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "5. 인재개발실이 커버할 수 있는 범위", 1)
    add_body(doc, "인재개발실은 사내 10년차 개발 역량(AI 도구 활용, 사내 앱 구축 경험)을 보유하고 있으며, 올거나이즈 플랫폼의 기술적 특성을 고려할 때 상당 부분을 내부에서 처리할 수 있습니다.")

    add_heading(doc, "5-1. 내부에서 충분히 커버 가능한 영역", 2)
    covers = [
        ("올거나이즈 플랫폼 기술 습득",
         "노코드 드래그앤드롭 도구이므로 인재개발실 내 개발 담당자가 1~2주면 마스터 가능. RAG 설정·워크플로우 구축·API 연동까지 내부 처리 가능"),
        ("기술 지원 업무 (견적의 400만원 항목)",
         "케이스별 솔루션 테스트, 프로토타입 구축, 기술 문제 해결 — 인재개발실 개발 역량으로 완전 대체 가능"),
        ("교안 기술 파트 작성",
         "LLM App Builder 실습 교안, RAG 설정 가이드, 워크플로우 다이어그램 등 기술 교안 내부에서 직접 작성 가능"),
        ("실습 보조·기술 멘토링",
         "워크숍 중 참가자 기술 질문 대응, 올거나이즈 환경 트러블슈팅"),
    ]
    for label, text in covers:
        add_icon_bullet(doc, "[O]", label, text, GREEN)

    add_heading(doc, "5-2. 준비 필요한 영역 (단기 학습으로 해결 가능)", 2)
    partial = [
        ("교안 기획 (워크숍 설계)",
         "ICE 프레임워크, 디자인씽킹 기반 Pain Point 발굴 설계는 퍼실리테이션 방법론 학습 필요 (온라인 강의 1~2일 분량)"),
        ("증권업 특화 케이스 발굴",
         "리서치·자산관리·IB 실무 이해가 필요하나, 인재개발실이 사내 현업 담당자와 협업하면 해결 가능"),
    ]
    for label, text in partial:
        add_icon_bullet(doc, "[△]", label, text, ORANGE)

    add_heading(doc, "5-3. 외부 지원이 필요할 수 있는 영역", 2)
    needs = [
        ("20명 대상 그룹 퍼실리테이션 진행",
         "교육학적 기법을 활용한 대규모 그룹 퍼실리테이션은 별도 스킬. 전문 퍼실리테이터 1명만 단기 조달 권장"),
        ("전사 AI 리터러시 커리큘럼 로드맵",
         "조직 전체 AI 역량 로드맵 설계는 교육공학·조직개발 전문성 영역 (이 워크숍과는 별개)"),
    ]
    for label, text in needs:
        add_icon_bullet(doc, "[○]", label, text, RED)

    add_heading(doc, "5-4. 비용 절감 시나리오", 2)
    add_body(doc, "인재개발실이 기술 지원과 교안 기술 파트를 내부에서 처리하면 기존 견적 대비 약 400~600만원 절감 가능:")
    add_bullet(doc, "기술 지원 400만원 → 인재개발실 내부 처리 (절감)")
    add_bullet(doc, "교안 설계 400만원 → 인재개발실이 기술 교안 작성 + 외부는 워크숍 설계만 → 약 150~200만원으로 축소")
    add_bullet(doc, "워크숍 진행 600만원 → 전문 퍼실리테이터만 외주 → 약 200~300만원으로 축소 (퍼실리테이터 1명 + 인재개발실 기술 지원)")
    add_body(doc, "예상 총액: 약 350~500만원 + VAT")

    # ══════════════════════════════════════════════════════════
    # 6. 내재화 가능성
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "6. 인재개발실 직원 내재화 가능성", 1)

    add_heading(doc, "6-1. 단계별 내재화 소요 기간", 2)
    add_table(doc,
        headers=["단계", "목표 수준", "소요 기간", "달성 방법"],
        rows=[
            ["1단계", "기본 조작", "6시간 워크숍", "문서 업로드, 간단 에이전트 생성"],
            ["2단계", "독립 활용", "워크숍 후 2~4주", "업무별 에이전트 스스로 설계·구축"],
            ["3단계", "타 직원 교육", "2~3개월 실무 경험", "반복 활용 후 내부 전파 가능"],
            ["4단계", "고도화·운영", "6개월 이상", "인재개발실 개발 담당자 지원 없이는 한계 존재"],
        ],
        col_widths=[2.0, 3.0, 3.5, 7.5],
    )

    add_heading(doc, "6-2. 내재화 성공의 핵심 요인", 2)
    add_bullet(doc, "워크숍 이후 올거나이즈 계정 유지 및 지속 실습 환경 확보 (가장 중요)")
    add_bullet(doc, "인재개발실 내 전담 기술 지원 담당자 존재 여부가 내재화 속도를 결정")
    add_bullet(doc, "월 1~2회 오피스아워(후속 세션) 운영 시 내재화 성공률 대폭 상승")
    add_bullet(doc, "실무에 즉시 적용 가능한 구체적 유스케이스 사전 선정 필요")

    add_heading(doc, "6-3. 현실적 기대치 (비개발 직원 20명 기준)", 2)
    add_table(doc,
        headers=["도달 수준", "예상 비율", "인원"],
        rows=[
            ["독립적으로 간단한 에이전트 구축 가능", "30~40%", "6~8명"],
            ["기존 에이전트를 수정·활용 가능", "40~50%", "8~10명"],
            ["워크숍 후 자발적 활용이 어려운 수준", "10~20%", "2~4명"],
        ],
        col_widths=[9.0, 3.5, 3.5],
    )

    # ══════════════════════════════════════════════════════════
    # 7. 시나리오 비교 및 권고
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "7. 시나리오 비교 및 최종 권고", 1)

    add_heading(doc, "7-1. 3가지 진행 시나리오 비교", 2)
    add_table(doc,
        headers=["시나리오", "예상 비용", "리스크", "특이사항"],
        rows=[
            ["A. 현 견적 그대로", "15,400천원", "높음",
             "기술지원 중복·레퍼런스 미검증·교안 소유권 불명확"],
            ["B. 협상 후 진행 (권장)", "7,700~9,900천원", "중간",
             "기술지원 삭제, 교안 소유권 귀속 조건 추가"],
            ["C. 인재개발실 주도 + FT 외주", "2,200~4,400천원", "낮음~중간",
             "인재개발실 업무 부하, 퍼실리테이션 품질 편차"],
        ],
        col_widths=[4.5, 3.5, 2.5, 5.5],
    )

    add_heading(doc, "7-2. 시나리오 B 협상 체크리스트", 2)
    nego = [
        "인일(Man-Day) 기반 상세 산출서 요청 — 투입 인력 수·등급·투입 일수 명시",
        "\"기술 지원\" 400만원 항목 삭제 또는 100만원 이하로 축소 협상 (인재개발실 내부 처리 + 올거나이즈 자체 온보딩 활용)",
        "\"그룹사 AI 리터러시 커리큘럼 자문\" 분리 — 이번 워크숍 교안 설계와 별도 계약으로 분리",
        "교안 및 실습 자료 재사용권을 하나증권에 귀속하도록 계약 조건 수정",
        "사후 지원 조건 명시 — 워크숍 후 최소 1개월 Q&A 채널 운영",
        "아리송컴퍼니 레퍼런스 제출 요청 — 유사 규모 금융권 워크숍 수행 실적",
        "올거나이즈 공식 파트너 여부 및 플랫폼 전문성 자격 확인",
    ]
    for i, item in enumerate(nego, 1):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(2)
        para.paragraph_format.left_indent  = Cm(0.3)
        r_n = para.add_run(f"{i}. ")
        set_font(r_n, size=10, bold=True, color=NAVY)
        r_t = para.add_run(item)
        set_font(r_t, size=10, color=GRAY)

    add_heading(doc, "7-3. 대안 검토", 2)
    add_bullet(doc,
        "올거나이즈 자체 교육 프로그램 확인: 이미 라이선스 비용을 내고 있다면 기본 교육이 포함되어 있을 수 있음. "
        "올거나이즈 담당 AM에게 교육 지원 범위 확인 후 견적 협상에 활용")
    add_bullet(doc,
        "인재개발실 주도 + 퍼실리테이터 1인 외주: 인재개발실이 기술 교안·실습 지원 전담, "
        "워크숍 진행만 전문 퍼실리테이터 1인 단기 계약. 예상 비용 200~400만원(VAT포함)")
    add_bullet(doc,
        "내부 파일럿 우선 진행: 인재개발실이 먼저 올거나이즈를 숙달 → 소규모(5명) 파일럿 내부 진행 "
        "→ 품질 확인 후 외부 교육 여부 결정. 리스크와 비용 동시 최소화")

    # ══════════════════════════════════════════════════════════
    # 8. 최종 결론
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "8. 최종 결론", 1)

    add_callout(doc,
        "핵심 결론 3가지",
        "① 기술 지원 400만원은 근거가 가장 약한 항목\n"
        "   → 올거나이즈 공식 온보딩과 중복, 인재개발실 내 개발 담당자가 1~2주 독학으로 완전 대체 가능\n\n"
        "② 아리송컴퍼니 레퍼런스가 사실상 없음\n"
        "   → 발주 전 유사 워크숍 수행 실적과 올거나이즈 공인 파트너 여부 반드시 서면 확인\n\n"
        "③ 협상 목표금액: 700~900만원 (VAT 포함)\n"
        "   → 기술지원 삭제·교안 소유권 귀속·사후지원 명시를 협상 카드로 활용\n"
        "   → 협상 불가 시 인재개발실 주도 + 퍼실리테이터 1인 외주로 200~400만원 수준에서 동등 품질 구현 가능",
        bg_hex=CALLOUT_BG)

    add_heading(doc, "의사결정 흐름", 2)
    add_table(doc,
        headers=["단계", "내용"],
        rows=[
            ["STEP 1", "올거나이즈 담당 AM에게 기본 교육 지원 범위 확인 (1~2일 내)"],
            ["STEP 2", "아리송컴퍼니에 상세 산출 근거 + 레퍼런스 서면 요청 (1주 내)"],
            ["STEP 3-A", "레퍼런스 충분 + 협상 성공 시: 700~900만원 수준으로 계약 (시나리오 B)"],
            ["STEP 3-B", "협상 결렬 또는 레퍼런스 불충분 시: 인재개발실 주도 내부 진행 (시나리오 C, 200~400만원)"],
        ],
        col_widths=[3.0, 13.0],
    )

    # Footer note
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_end.paragraph_format.space_before = Pt(20)
    r_end = p_end.add_run(
        "본 보고서는 2026년 4월 12일 기준으로 작성된 내부 검토 자료입니다.  |  외부 배포 금지  |  하나증권 인재개발실"
    )
    set_font(r_end, size=9, color=LGRAY)

    # ── Save ──────────────────────────────────────────────────
    out_path = r"c:\dev\hrd-workspace\하나증권_AX Sprint_견적분석_보고서.docx"
    doc.save(out_path)
    print(f"Word 파일 생성 완료: {out_path}")


if __name__ == "__main__":
    build()
