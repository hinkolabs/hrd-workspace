"""
ICE 프레임워크 가이드 Word 파일 생성 스크립트
실행: python generate_ice_guide_word.py
출력: ice_framework_guide.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

OUTPUT_FILE = "ice_framework_guide.docx"

# ── 색상 상수 ──────────────────────────────────────────────
C_INDIGO     = RGBColor(0x4F, 0x46, 0xE5)
C_ROSE       = RGBColor(0xE1, 0x1D, 0x48)
C_EMERALD    = RGBColor(0x05, 0x96, 0x6D)
C_AMBER      = RGBColor(0xD9, 0x77, 0x06)
C_GRAY       = RGBColor(0x6B, 0x72, 0x80)
C_DARK       = RGBColor(0x11, 0x18, 0x27)
C_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
C_HEADER_BG  = RGBColor(0x4F, 0x46, 0xE5)
C_ROW_LIGHT  = RGBColor(0xF5, 0xF3, 0xFF)

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)

def set_cell_border(cell, border_color="C7D2FE"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), border_color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1, color=None):
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
    p = doc.add_paragraph(style=style_map.get(level, "Heading 1"))
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt({1: 16, 2: 13, 3: 11}.get(level, 13))
    if color:
        run.font.color.rgb = color
    return p

def add_body(doc, text, bold=False, color=None, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(doc, text, color=None):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = color
    return p

def build_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = C_WHITE
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, C_HEADER_BG)

    # Data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = C_ROW_LIGHT if ri % 2 == 0 else RGBColor(0xFF, 0xFF, 0xFF)
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            if ci == 0:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = C_INDIGO
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_bg(cell, bg)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)

    return table


def main():
    doc = Document()

    # ── 문서 기본 설정 ──────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Normal 스타일 기본 폰트
    doc.styles["Normal"].font.name = "맑은 고딕"
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.styles["Normal"].element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    # ── 표지 ────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\nICE 프레임워크")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = C_INDIGO

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("AI 에이전트 개발 우선순위 선정 가이드")
    r2.font.size = Pt(15)
    r2.font.color.rgb = C_GRAY

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"인재개발실  |  {datetime.date.today().strftime('%Y년 %m월')}")
    r3.font.size = Pt(11)
    r3.font.color.rgb = C_GRAY

    doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # 1. ICE 프레임워크란?
    # ══════════════════════════════════════════════════════
    add_heading(doc, "1. ICE 프레임워크란?", 1, C_INDIGO)
    add_body(doc,
        "ICE는 여러 아이디어나 Pain Point 중 어떤 것을 먼저 개발할 것인가를 "
        "팀이 주관적 느낌이 아닌 숫자 근거로 합의하게 해주는 우선순위 결정 도구입니다.",
        size=11
    )
    doc.add_paragraph()

    # ICE 공식
    formula_p = doc.add_paragraph()
    formula_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_p.add_run("ICE Score  =  I (Impact)  ×  C (Confidence)  ×  E (Ease)").bold = True
    formula_p.runs[0].font.size = Pt(14)
    formula_p.runs[0].font.color.rgb = C_INDIGO

    doc.add_paragraph()
    add_body(doc, "각 항목을 1~10점으로 평가하여 최대 1,000점까지 산출됩니다.", color=C_GRAY)
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════
    # 2. 채점 기준
    # ══════════════════════════════════════════════════════
    add_heading(doc, "2. 채점 기준", 1, C_INDIGO)

    axes = [
        {
            "title": "I · Impact (임팩트)",
            "color": C_ROSE,
            "desc": "이 문제를 해결했을 때 업무에 얼마나 큰 효과가 있나요?",
            "rows": [
                ("9~10점", "주 5시간 이상 절감, 매출·리스크에 직접 영향"),
                ("7~8점",  "주 2~4시간 절감, 오류율 대폭 감소"),
                ("4~6점",  "주 1시간 이하 절감, 편의성 개선"),
                ("1~3점",  "체감 효과 미미, 있으면 좋은 수준"),
            ]
        },
        {
            "title": "C · Confidence (확신도)",
            "color": C_INDIGO,
            "desc": "이 문제가 실제로 존재하고 AI로 해결 가능하다고 얼마나 확신하나요?",
            "rows": [
                ("9~10점", "데이터·수치로 문제 존재 입증 가능"),
                ("7~8점",  "팀원 과반이 동일 문제 경험"),
                ("4~6점",  "일부만 경험, 빈도 불확실"),
                ("1~3점",  "추측 수준, 근거 데이터 없음"),
            ]
        },
        {
            "title": "E · Ease (실현 용이성)",
            "color": C_EMERALD,
            "desc": "올거나이즈(Alli)로 구현하는 것이 얼마나 쉬운가요?",
            "rows": [
                ("9~10점", "문서 업로드 + Q&A 에이전트로 해결 (1~2시간)"),
                ("7~8점",  "노드 3~5개 워크플로우로 해결 (반나절)"),
                ("4~6점",  "외부 시스템 API 연동 필요 (1~2일)"),
                ("1~3점",  "핵심 데이터가 외부 시스템에 있어 연동 난이도 높음"),
            ]
        },
    ]

    for axis in axes:
        add_heading(doc, axis["title"], 2, axis["color"])
        add_body(doc, axis["desc"], color=C_GRAY)
        build_table(doc, ["점수", "기준"], axis["rows"], col_widths=[2.5, 11])
        doc.add_paragraph()

    # ══════════════════════════════════════════════════════
    # 3. 채점 예시
    # ══════════════════════════════════════════════════════
    add_heading(doc, "3. 채점 예시", 1, C_INDIGO)

    example_rows = [
        ("사규/규정집 질문 자동 응답",  "8", "9", "9", "648", "즉시 실행"),
        ("신입사원 온보딩 Q&A 봇",     "7", "8", "9", "504", "즉시 실행"),
        ("리서치 보고서 요약 자동화",   "9", "7", "7", "441", "우선 검토"),
        ("고객 계약서 핵심 조항 추출",  "8", "8", "5", "320", "우선 검토"),
        ("ERP 데이터 자동 리포팅",     "9", "6", "2", "108", "장기 과제"),
    ]

    build_table(
        doc,
        ["Pain Point", "I", "C", "E", "ICE Score", "판정"],
        example_rows,
        col_widths=[6.5, 1.5, 1.5, 1.5, 2.5, 2.5]
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════
    # 4. 점수 해석 가이드
    # ══════════════════════════════════════════════════════
    add_heading(doc, "4. 점수 해석 가이드", 1, C_INDIGO)

    priority_rows = [
        ("500점 이상", "즉시 실행",   "워크숍 당일 프로토타입 제작 대상"),
        ("200~499점", "우선 검토",   "다음 스프린트 후보 항목"),
        ("100~199점", "장기 과제",   "추가 검토 후 로드맵에 편입"),
        ("100점 미만", "보류",        "이번 워크숍 범위 외로 분리"),
    ]

    build_table(
        doc,
        ["점수 범위", "판정", "설명"],
        priority_rows,
        col_widths=[3, 3, 10]
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════
    # 5. Impact × Ease 사분면 매트릭스
    # ══════════════════════════════════════════════════════
    add_heading(doc, "5. Impact × Ease 사분면 매트릭스", 1, C_INDIGO)

    add_body(doc,
        "ICE Score가 높더라도 Ease가 낮은 항목(전략 과제)은 당장 워크숍에서 만들기 어렵습니다.\n"
        "Quick Win 영역을 워크숍 첫 번째 구현 대상으로 선정하고, 전략 과제는 별도 프로젝트로 추진하는 것을 권장합니다.",
        size=10.5
    )
    doc.add_paragraph()

    quadrant_rows = [
        ("Impact 높음\nEase 높음",  "Quick Win ★",  "워크숍 당일 제작 가능, 최우선"),
        ("Impact 높음\nEase 낮음",  "전략 과제",     "중장기 로드맵 편입"),
        ("Impact 낮음\nEase 높음",  "패스",          "우선순위 낮음"),
        ("Impact 낮음\nEase 낮음",  "장기 검토",     "현재는 보류"),
    ]

    build_table(
        doc,
        ["조건", "사분면", "대응"],
        quadrant_rows,
        col_widths=[4, 3, 9]
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════
    # 6. 워크숍 활용 흐름
    # ══════════════════════════════════════════════════════
    add_heading(doc, "6. 워크숍 활용 흐름", 1, C_INDIGO)

    flow_rows = [
        ("1", "사전 준비\n(1~2주 전)",
         "• Pain Point 설문 링크 생성\n• 참가자 20명에게 설문 배포\n• 응답 수집 및 모니터링"),
        ("2", "Session 1 (60분)",
         "• ICE 채점 기준 설명 (5분)\n• Pain Point 목록 공유 (10분)\n• 팀별 I·C·E 점수 부여 (30분)\n• 상위 항목 합의 선정 (15분)"),
        ("3", "Session 2 (60분)",
         "• 선정된 항목 프로토타입 설계\n• 올거나이즈 구현 실습\n• 팀별 결과 발표 및 피드백"),
        ("4", "결과 활용",
         "• ICE Score 결과 CSV 저장\n• Quick Win 항목 즉시 개발 착수\n• 로드맵에 중장기 과제 편입"),
    ]

    build_table(
        doc,
        ["단계", "구분", "세부 내용"],
        flow_rows,
        col_widths=[1.5, 4, 10.5]
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════
    # 7. 사전 설문 기본 질문 (참고용)
    # ══════════════════════════════════════════════════════
    add_heading(doc, "7. 사전 설문 기본 질문 (참고용)", 1, C_INDIGO)

    add_body(doc, "아래 3가지 기본 질문을 바탕으로 설문을 구성하세요. 필요에 따라 질문을 추가·변경할 수 있습니다.")
    doc.add_paragraph()

    for i, q in enumerate([
        "현재 업무에서 가장 반복적이고 시간이 많이 걸리는 작업은 무엇인가요?",
        "AI가 자동화해줬으면 하는 업무를 구체적으로 적어주세요 (최대 3가지)",
        "위 업무를 처리하는 데 주당 몇 시간 정도 소요되나요?\n   (1시간 미만 / 1~3시간 / 3~5시간 / 5~10시간 / 10시간 이상)"
    ], 1):
        p = doc.add_paragraph()
        run = p.add_run(f"Q{i}. {q}")
        run.font.size = Pt(10.5)

    doc.add_paragraph()

    # ── 푸터 메모 ─────────────────────────────────────────
    doc.add_page_break()
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_end = p_end.add_run(
        f"본 문서는 HRD Workspace ICE 우선순위 채점 시스템에서 생성되었습니다.\n"
        f"생성일: {datetime.date.today().strftime('%Y년 %m월 %d일')}  |  인재개발실"
    )
    r_end.font.size = Pt(9)
    r_end.font.color.rgb = C_GRAY

    # ── 저장 ──────────────────────────────────────────────
    doc.save(OUTPUT_FILE)
    print(f"\n[완료] '{OUTPUT_FILE}' 파일이 생성되었습니다.")
    print(f"경로: {os.path.abspath(OUTPUT_FILE)}\n")


if __name__ == "__main__":
    main()
