import sys
sys.stdout.reconfigure(encoding='utf-8')
sys.stderr.reconfigure(encoding='utf-8')

import docx

doc = docx.Document(r'C:\Users\jay\Downloads\하나증권_신입연수_AI교육_강사용_진행큐시트.docx')

print(f'총 단락 수: {len(doc.paragraphs)}')
print(f'총 테이블 수: {len(doc.tables)}')

print('\n========== 단락 전체 ==========')
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        style = para.style.name if para.style else ''
        print(f'[{i:3d}][{style}] {text}')

print('\n========== 테이블 전체 ==========')
for t_idx, table in enumerate(doc.tables):
    print(f'\n--- 테이블 {t_idx+1} ({len(table.rows)}행 x {len(table.columns)}열) ---')
    for r_idx, row in enumerate(table.rows):
        cells = [cell.text.strip().replace('\n', ' / ') for cell in row.cells]
        # 중복 셀 제거 (병합된 셀)
        seen = []
        deduped = []
        for c in cells:
            if c not in seen:
                seen.append(c)
                deduped.append(c)
        line = ' | '.join(deduped)
        if line.strip():
            print(f'  행{r_idx+1:3d}: {line[:200]}')
